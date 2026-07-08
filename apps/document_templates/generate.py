import re
import uuid
from datetime import date
from io import BytesIO

from django.core.files.base import ContentFile
from docx import Document as DocxDocument

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def _fmt_date(value):
    return value.strftime("%d.%m.%Y") if value else ""


def _fmt_money(value):
    if value is None:
        return ""
    return f"{int(round(value)):,}".replace(",", " ")


def build_context(project, employee_contract=None, extra_fields=None):
    context = {
        "client_name": project.client_name or "",
        "client_contact": project.client_contact or "",
        "project_name": project.name,
        "address": project.address or "",
        "budget": _fmt_money(project.budget),
        "paid_amount": _fmt_money(project.paid_amount),
        "start_date": _fmt_date(project.start_date),
        "deadline": _fmt_date(project.deadline),
        "today": _fmt_date(date.today()),
    }
    if employee_contract is not None:
        context.update({
            "employee_name": employee_contract.user.full_name,
            "employee_amount": _fmt_money(employee_contract.amount),
            "employee_advance": _fmt_money(employee_contract.advance),
            "employee_paid": _fmt_money(employee_contract.paid),
        })
    if extra_fields:
        context.update(extra_fields)
    return context


def render_content(content, context):
    def replace(match):
        key = match.group(1)
        return str(context.get(key, match.group(0)))
    return PLACEHOLDER_RE.sub(replace, content)


def render_docx(title, rendered_text):
    doc = DocxDocument()
    doc.add_heading(title, level=1)
    for line in rendered_text.splitlines():
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    filename = f"{uuid.uuid4().hex}.docx"
    return filename, ContentFile(buffer.read(), name=filename)
