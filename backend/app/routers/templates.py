import os
import re
import uuid
from datetime import datetime

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.config import settings
from app.database import get_db
from app.models.user import User, UserRole
from app.models.project import Project
from app.models.finance import Contract, EmployeeContract
from app.models.document import Document, AuditLog
from app.models.template import DocumentTemplate
from app.schemas.template import (
    TemplateCreate, TemplateUpdate, TemplateResponse, TemplateGenerateRequest,
)
from app.utils.dependencies import get_current_active_user, require_roles

router = APIRouter(prefix="/templates", tags=["templates"])

PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


@router.get("", response_model=list[TemplateResponse])
async def list_templates(
    template_type: str | None = None,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    query = select(DocumentTemplate)
    if template_type:
        query = query.where(DocumentTemplate.template_type == template_type)
    result = await db.execute(query.order_by(DocumentTemplate.created_at.desc()))
    return result.scalars().all()


@router.post("", response_model=TemplateResponse, status_code=201)
async def create_template(
    data: TemplateCreate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_roles(UserRole.admin, UserRole.manager, UserRole.gip)),
):
    template = DocumentTemplate(**data.model_dump(), created_by=current_user.id)
    db.add(template)
    await db.commit()
    await db.refresh(template)
    return template


@router.get("/{template_id}", response_model=TemplateResponse)
async def get_template(
    template_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    return template


@router.put("/{template_id}", response_model=TemplateResponse)
async def update_template(
    template_id: str,
    data: TemplateUpdate,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_roles(UserRole.admin, UserRole.manager, UserRole.gip)),
):
    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    for field, value in data.model_dump(exclude_none=True).items():
        setattr(template, field, value)
    await db.commit()
    await db.refresh(template)
    return template


@router.delete("/{template_id}")
async def delete_template(
    template_id: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(require_roles(UserRole.admin, UserRole.manager)),
):
    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")
    await db.delete(template)
    await db.commit()
    return {"message": "Template deleted"}


async def _build_context(db: AsyncSession, data: TemplateGenerateRequest) -> dict[str, str]:
    result = await db.execute(select(Project).where(Project.id == data.project_id))
    project = result.scalar_one_or_none()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    def fmt_date(d):
        return d.strftime("%d.%m.%Y") if d else ""

    def fmt_money(v):
        return f"{v:,.0f}".replace(",", " ") if v is not None else ""

    context: dict[str, str] = {
        "today": datetime.utcnow().strftime("%d.%m.%Y"),
        "project_name": project.name or "",
        "project_description": project.description or "",
        "client_name": project.client_name or "",
        "client_contact": project.client_contact or "",
        "address": project.address or "",
        "project_stage": project.stage.value if project.stage else "",
        "project_status": project.status.value if project.status else "",
        "start_date": fmt_date(project.start_date),
        "deadline": fmt_date(project.deadline),
        "budget": fmt_money(project.budget),
        "paid_amount": fmt_money(project.paid_amount),
        "amount": fmt_money(project.budget),
    }

    if data.contract_id:
        result = await db.execute(select(Contract).where(Contract.id == data.contract_id))
        contract = result.scalar_one_or_none()
        if contract:
            context.update({
                "contract_number": contract.contract_number or "",
                "contract_amount": fmt_money(contract.amount),
                "contract_date": fmt_date(contract.signed_date),
                "contract_deadline": fmt_date(contract.deadline),
                "client_name": contract.client_name or context["client_name"],
                "amount": fmt_money(contract.amount),
            })

    if data.employee_id:
        result = await db.execute(select(User).where(User.id == data.employee_id))
        employee = result.scalar_one_or_none()
        if employee:
            context.update({
                "employee_name": employee.full_name or "",
                "employee_email": employee.email or "",
                "employee_phone": employee.phone or "",
            })
        ec_result = await db.execute(
            select(EmployeeContract).where(
                EmployeeContract.user_id == data.employee_id,
                EmployeeContract.project_id == data.project_id,
            )
        )
        ec = ec_result.scalars().first()
        if ec:
            context.update({
                "employee_contract_amount": fmt_money(ec.amount),
                "employee_advance": fmt_money(ec.advance),
                "employee_paid": fmt_money(ec.paid),
                "employee_balance": fmt_money(ec.amount - ec.paid),
            })

    context.update({k: str(v) for k, v in data.extra_fields.items()})
    return context


def _render_docx(text: str, out_path: str, title: str):
    from docx import Document as DocxDocument

    docx = DocxDocument()
    docx.add_heading(title, level=1)
    for line in text.split("\n"):
        docx.add_paragraph(line)
    docx.save(out_path)


@router.post("/{template_id}/generate")
async def generate_document(
    template_id: str,
    data: TemplateGenerateRequest,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(get_current_active_user),
):
    result = await db.execute(select(DocumentTemplate).where(DocumentTemplate.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="Template not found")

    context = await _build_context(db, data)
    rendered = PLACEHOLDER_RE.sub(lambda m: context.get(m.group(1), ""), template.content)

    out_dir = os.path.join(settings.UPLOAD_DIR, "contracts")
    os.makedirs(out_dir, exist_ok=True)
    filename = f"{template.template_type.value}_{uuid.uuid4().hex[:8]}.docx"
    out_path = os.path.join(out_dir, filename)
    doc_title = f"{template.name} — {context.get('project_name', '')}"
    _render_docx(rendered, out_path, doc_title)
    file_size = os.path.getsize(out_path)

    document_id = None
    if data.save_as_document:
        document = Document(
            name=f"{doc_title}.docx",
            doc_type=template.template_type.value,
            project_id=data.project_id,
            uploaded_by=current_user.id,
            file_path=out_path,
            file_size=file_size,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        db.add(document)
        await db.flush()
        document_id = document.id
        db.add(AuditLog(
            entity_type="document", entity_id=document.id,
            action="generate_from_template", user_id=current_user.id,
            details={"template_id": template_id, "template_name": template.name},
        ))
        await db.commit()

    return {
        "message": "Document generated",
        "document_id": document_id,
        "file_path": out_path,
        "rendered_text": rendered,
        "download_url": f"/api/templates/generated/{filename}",
    }


@router.get("/generated/{filename}")
async def download_generated(
    filename: str,
    current_user: User = Depends(get_current_active_user),
):
    safe_name = os.path.basename(filename)
    path = os.path.join(settings.UPLOAD_DIR, "contracts", safe_name)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path, filename=safe_name)
